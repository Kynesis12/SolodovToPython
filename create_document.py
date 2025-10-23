#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os

def create_gost_document():
    """Создает документ Word с оформлением по российскому ГОСТу"""
    
    # Создаем новый документ
    doc = Document()
    
    # Настройка стилей по ГОСТу
    setup_gost_styles(doc)
    
    # Добавляем титульную страницу
    add_title_page(doc)
    
    # Добавляем оглавление
    add_table_of_contents(doc)
    
    # Читаем и добавляем все лекции
    lectures_dir = "/media/maskpov/Povarov_IVT44u/СЕМЕСТР 3/Python/SolodovToPython/питон"
    lecture_files = [
        "00_Введение_и_титульная_часть_Python.txt",
        "01_Лекция_1_Процесс_создания_ПО_Python.txt", 
        "02_Лекция_2_Среда_разработки_Python.txt",
        "03_Лекция_3_Типы_данных_Python.txt",
        "04_Лекция_4_Операторы_Python.txt",
        "05_Лекция_5_Инструкции_управления_Python.txt",
        "06_Лекция_6_Циклы_Python.txt",
        "07_Лекция_7_Массивы_Python.txt",
        "08_Лекция_8_Символы_и_строки_Python.txt",
        "09_Лекция_9_Классы_объекты_и_методы_Python.txt",
        "10_Лекция_10_Исключения_Python.txt",
        "11_Лекция_11_Приложение_под_ОС_Windows_Python.txt",
        "12_Лекция_12_Графика_в_Python.txt"
    ]
    
    for i, filename in enumerate(lecture_files):
        filepath = os.path.join(lectures_dir, filename)
        if os.path.exists(filepath):
            add_lecture_content(doc, filepath, i)
    
    # Сохраняем документ
    output_path = "/media/maskpov/Povarov_IVT44u/СЕМЕСТР 3/Python/SolodovToPython/Учебное_пособие_Python_ГОСТ.docx"
    doc.save(output_path)
    print(f"Документ сохранен: {output_path}")

def setup_gost_styles(doc):
    """Настройка стилей по ГОСТу"""
    
    # Основной стиль текста
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(14)
    normal_paragraph = normal_style.paragraph_format
    normal_paragraph.line_spacing = 1.5
    normal_paragraph.first_line_indent = Cm(1.27)
    normal_paragraph.space_after = Pt(6)
    
    # Стиль заголовка 1
    heading1_style = doc.styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Times New Roman'
    heading1_font.size = Pt(16)
    heading1_font.bold = True
    heading1_paragraph = heading1_style.paragraph_format
    heading1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1_paragraph.space_before = Pt(12)
    heading1_paragraph.space_after = Pt(12)
    
    # Стиль заголовка 2
    heading2_style = doc.styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Times New Roman'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2_paragraph = heading2_style.paragraph_format
    heading2_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2_paragraph.space_before = Pt(12)
    heading2_paragraph.space_after = Pt(6)
    
    # Стиль заголовка 3
    heading3_style = doc.styles['Heading 3']
    heading3_font = heading3_style.font
    heading3_font.name = 'Times New Roman'
    heading3_font.size = Pt(14)
    heading3_font.bold = True
    heading3_paragraph = heading3_style.paragraph_format
    heading3_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3_paragraph.space_before = Pt(6)
    heading3_paragraph.space_after = Pt(6)

def add_title_page(doc):
    """Добавляет титульную страницу"""
    
    # Заголовок министерства
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("МИНИСТЕРСТВО ЦИФРОВОГО РАЗВИТИЯ, СВЯЗИ И\nМАССОВЫХ КОММУНИКАЦИЙ РФ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Пустые строки
    for _ in range(2):
        doc.add_paragraph()
    
    # Название университета
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Федеральное государственное бюджетное образовательное учреждение\nвысшего образования\n«ПОВОЛЖСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ\nТЕЛЕКОММУНИКАЦИЙ И ИНФОРМАТИКИ»")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Пустые строки
    for _ in range(3):
        doc.add_paragraph()
    
    # Кафедра
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Кафедра информатики и вычислительной техники")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Пустые строки
    for _ in range(4):
        doc.add_paragraph()
    
    # Авторы
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Т.А. КОВАЛЕНКО, А.Г. СОЛОДОВ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Пустые строки
    for _ in range(2):
        doc.add_paragraph()
    
    # Название пособия
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Вычислительная техника и языки программирования\n\nЧасть 1 - Python Edition\n\n\nУЧЕБНОЕ ПОСОБИЕ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Пустые строки
    for _ in range(8):
        doc.add_paragraph()
    
    # Год и место
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Самара, 2021")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Разрыв страницы
    doc.add_page_break()

def add_table_of_contents(doc):
    """Добавляет оглавление"""
    
    # Заголовок оглавления
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("СОДЕРЖАНИЕ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Содержание
    contents = [
        ("Введение", "3"),
        ("Лекция 1. Процесс создания программного обеспечения", "4"),
        ("Лекция 2. Среда разработки Python", "8"),
        ("Лекция 3. Типы данных в Python", "12"),
        ("Лекция 4. Операторы в Python", "16"),
        ("Лекция 5. Инструкции управления", "20"),
        ("Лекция 6. Циклы в Python", "24"),
        ("Лекция 7. Массивы (списки) в Python", "28"),
        ("Лекция 8. Символы и строки в Python", "32"),
        ("Лекция 9. Классы, объекты и методы в Python", "36"),
        ("Лекция 10. Исключения в Python", "40"),
        ("Лекция 11. Приложения под ОС Windows в Python", "44"),
        ("Лекция 12. Графика в Python", "48")
    ]
    
    for title, page in contents:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 1.0
        
        run1 = p.add_run(title)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(14)
        
        # Добавляем точки до номера страницы
        dots = "." * (50 - len(title))
        run2 = p.add_run(dots)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(14)
        
        run3 = p.add_run(page)
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(14)
    
    # Разрыв страницы
    doc.add_page_break()

def add_lecture_content(doc, filepath, lecture_num):
    """Добавляет содержимое лекции в документ"""
    
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        
        # Обрабатываем каждую строку
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line:
                continue
            
            # Определяем тип строки и применяем соответствующий стиль
            if line.startswith('Лекция') and ':' in line:
                # Заголовок лекции
                p = doc.add_paragraph(line)
                p.style = 'Heading 1'
                
            elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.') or line.startswith('5.'):
                # Нумерованный список
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                p.paragraph_format.first_line_indent = Cm(0.5)
                
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # Маркированный список
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                p.paragraph_format.first_line_indent = Cm(0.5)
                
            elif line.startswith('Пример') and ':' in line:
                # Заголовок примера
                p = doc.add_paragraph(line)
                p.style = 'Heading 3'
                
            elif line.startswith('def ') or line.startswith('class ') or line.startswith('import ') or line.startswith('from '):
                # Код Python
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(12)
                
            elif line.startswith('#') or line.startswith('//'):
                # Комментарии в коде
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(12)
                run.font.italic = True
                
            elif line.startswith('    ') and (line.strip().startswith('print') or line.strip().startswith('if') or line.strip().startswith('for') or line.strip().startswith('while')):
                # Отступы в коде
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(12)
                
            else:
                # Обычный текст
                p = doc.add_paragraph(line)
                p.style = 'Normal'
        
        # Добавляем разрыв страницы между лекциями (кроме последней)
        if lecture_num < 12:
            doc.add_page_break()
            
    except Exception as e:
        print(f"Ошибка при обработке файла {filepath}: {e}")

if __name__ == "__main__":
    create_gost_document()
    print("Документ успешно создан!")
